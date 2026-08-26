"""Genera los manuales en Word de Papelería de Sol.

Se usa el preajuste compacto de guía de referencia y una portada editorial
con los colores de la marca como excepción visual documentada.
"""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


RAIZ = Path(__file__).resolve().parents[1]
CARPETA_SALIDA = RAIZ / "documentacion" / "manuales"
LOGO = RAIZ / "public" / "marca" / "logo-papeleria-de-sol-transparente.png"

AZUL_GUIA = RGBColor(0x2E, 0x74, 0xB5)
AZUL_OSCURO_GUIA = RGBColor(0x1F, 0x4D, 0x78)
TINTA = RGBColor(0x25, 0x24, 0x34)
LILA_SOL = RGBColor(0x9B, 0x58, 0xC5)
LILA_SUAVE = "F6F0FA"
GRIS = RGBColor(0x64, 0x63, 0x73)


def establecer_fuente(run, nombre="Calibri", tamano=11, color=None, negrita=None, cursiva=None):
    run.font.name = nombre
    run._element.rPr.rFonts.set(qn("w:ascii"), nombre)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), nombre)
    run.font.size = Pt(tamano)
    if color:
        run.font.color.rgb = color
    if negrita is not None:
        run.bold = negrita
    if cursiva is not None:
        run.italic = cursiva


def sombrear_parrafo(parrafo, relleno=LILA_SUAVE):
    propiedades = parrafo._p.get_or_add_pPr()
    sombreado = OxmlElement("w:shd")
    sombreado.set(qn("w:fill"), relleno)
    propiedades.append(sombreado)


def agregar_borde_inferior(parrafo, color="9B58C5"):
    propiedades = parrafo._p.get_or_add_pPr()
    bordes = OxmlElement("w:pBdr")
    borde = OxmlElement("w:bottom")
    borde.set(qn("w:val"), "single")
    borde.set(qn("w:sz"), "12")
    borde.set(qn("w:space"), "6")
    borde.set(qn("w:color"), color)
    bordes.append(borde)
    propiedades.append(bordes)


def agregar_campo_pagina(parrafo):
    inicio = OxmlElement("w:fldChar")
    inicio.set(qn("w:fldCharType"), "begin")
    instruccion = OxmlElement("w:instrText")
    instruccion.set(qn("xml:space"), "preserve")
    instruccion.text = " PAGE "
    separado = OxmlElement("w:fldChar")
    separado.set(qn("w:fldCharType"), "separate")
    texto = OxmlElement("w:t")
    texto.text = "1"
    fin = OxmlElement("w:fldChar")
    fin.set(qn("w:fldCharType"), "end")
    run = parrafo.add_run()
    run._r.append(inicio)
    run._r.append(instruccion)
    run._r.append(separado)
    run._r.append(texto)
    run._r.append(fin)
    establecer_fuente(run, tamano=9, color=GRIS)


def configurar_estilos(documento):
    """Aplica los tokens del preajuste compact_reference_guide.

    Excepción nombrada: los elementos de marca usan lila (#9B58C5) en vez del
    azul del preajuste para mantener la identidad de Papelería de Sol.
    """
    estilos = documento.styles
    normal = estilos["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for nombre, tamano, color, antes, despues in [
        ("Heading 1", 16, AZUL_GUIA, 18, 10),
        ("Heading 2", 13, AZUL_GUIA, 14, 7),
        ("Heading 3", 12, AZUL_OSCURO_GUIA, 10, 5),
    ]:
        estilo = estilos[nombre]
        estilo.font.name = "Calibri"
        estilo._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        estilo._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        estilo.font.size = Pt(tamano)
        estilo.font.color.rgb = color
        estilo.font.bold = True
        estilo.paragraph_format.space_before = Pt(antes)
        estilo.paragraph_format.space_after = Pt(despues)
        estilo.paragraph_format.keep_with_next = True

    for nombre in ("List Bullet", "List Number"):
        estilo = estilos[nombre]
        estilo.font.name = "Calibri"
        estilo._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        estilo._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        estilo.font.size = Pt(11)
        estilo.paragraph_format.left_indent = Inches(0.375)
        estilo.paragraph_format.first_line_indent = Inches(-0.188)
        estilo.paragraph_format.space_after = Pt(4)
        estilo.paragraph_format.line_spacing = 1.25

    estilo_nota = estilos.add_style("Nota importante", WD_STYLE_TYPE.PARAGRAPH)
    estilo_nota.base_style = normal
    estilo_nota.paragraph_format.space_before = Pt(6)
    estilo_nota.paragraph_format.space_after = Pt(8)
    estilo_nota.paragraph_format.left_indent = Inches(0.18)
    estilo_nota.paragraph_format.right_indent = Inches(0.18)


def configurar_documento(titulo_corto):
    documento = Document()
    seccion = documento.sections[0]
    seccion.page_width = Inches(8.5)
    seccion.page_height = Inches(11)
    seccion.top_margin = Inches(1)
    seccion.right_margin = Inches(1)
    seccion.bottom_margin = Inches(1)
    seccion.left_margin = Inches(1)
    seccion.header_distance = Inches(0.492)
    seccion.footer_distance = Inches(0.492)
    configurar_estilos(documento)

    encabezado = seccion.header.paragraphs[0]
    encabezado.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = encabezado.add_run(f"Papelería de Sol | {titulo_corto}")
    establecer_fuente(run, tamano=9, color=GRIS, negrita=True)

    pie = seccion.footer.paragraphs[0]
    pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    texto = pie.add_run("Papelería de Sol | Página ")
    establecer_fuente(texto, tamano=9, color=GRIS)
    agregar_campo_pagina(pie)
    return documento


def agregar_portada(documento, titulo, subtitulo):
    """Patrón editorial_cover: portada centrada, aireada y con identidad propia."""
    documento.add_paragraph().paragraph_format.space_after = Pt(42)
    if LOGO.exists():
        parrafo_logo = documento.add_paragraph()
        parrafo_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        imagen = parrafo_logo.add_run().add_picture(str(LOGO), width=Inches(4.2))
        imagen._inline.docPr.set("descr", "Logotipo de Papelería de Sol")

    etiqueta = documento.add_paragraph()
    etiqueta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    etiqueta.paragraph_format.space_before = Pt(22)
    etiqueta.paragraph_format.space_after = Pt(16)
    run_etiqueta = etiqueta.add_run("GUÍA DE USO")
    establecer_fuente(run_etiqueta, tamano=11, color=LILA_SOL, negrita=True)

    p_titulo = documento.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_titulo.paragraph_format.space_after = Pt(10)
    run_titulo = p_titulo.add_run(titulo)
    establecer_fuente(run_titulo, tamano=27, color=TINTA, negrita=True)

    p_subtitulo = documento.add_paragraph()
    p_subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_subtitulo.paragraph_format.space_after = Pt(42)
    run_subtitulo = p_subtitulo.add_run(subtitulo)
    establecer_fuente(run_subtitulo, tamano=13, color=GRIS)
    agregar_borde_inferior(p_subtitulo)

    p_fecha = documento.add_paragraph()
    p_fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_fecha = p_fecha.add_run(f"Versión inicial | {date.today().strftime('%d/%m/%Y')}")
    establecer_fuente(run_fecha, tamano=10, color=LILA_SOL, negrita=True)

    documento.add_page_break()


def agregar_titulo(documento, texto):
    return documento.add_paragraph(texto, style="Heading 1")


def agregar_subtitulo(documento, texto):
    return documento.add_paragraph(texto, style="Heading 2")


def agregar_parrafo(documento, texto, negrita_inicial=None):
    parrafo = documento.add_paragraph()
    if negrita_inicial and texto.startswith(negrita_inicial):
        primero = parrafo.add_run(negrita_inicial)
        establecer_fuente(primero, negrita=True)
        resto = parrafo.add_run(texto[len(negrita_inicial):])
        establecer_fuente(resto)
    else:
        run = parrafo.add_run(texto)
        establecer_fuente(run)
    return parrafo


def agregar_nota(documento, texto):
    parrafo = documento.add_paragraph(style="Nota importante")
    sombrear_parrafo(parrafo)
    etiqueta = parrafo.add_run("Importante. ")
    establecer_fuente(etiqueta, color=LILA_SOL, negrita=True)
    cuerpo = parrafo.add_run(texto)
    establecer_fuente(cuerpo)
    return parrafo


def agregar_lista(documento, elementos, estilo="List Bullet"):
    for elemento in elementos:
        parrafo = documento.add_paragraph(style=estilo)
        run = parrafo.add_run(elemento)
        establecer_fuente(run)


def crear_manual_clientes():
    documento = configurar_documento("Manual para clientes")
    agregar_portada(
        documento,
        "Manual de uso para clientes",
        "Cómo explorar productos y realizar una consulta por WhatsApp.",
    )

    agregar_titulo(documento, "1. Propósito de este manual")
    agregar_parrafo(
        documento,
        "Esta guía explica cómo recorrer Papelería de Sol, armar una selección de productos y continuar la consulta por WhatsApp. La tienda muestra el catálogo, pero la disponibilidad, el pago y la entrega se coordinan directamente por ese medio.",
    )

    agregar_titulo(documento, "2. Recorrer la tienda")
    agregar_subtitulo(documento, "Usar el menú")
    agregar_parrafo(documento, "El menú superior permite acceder a Inicio, Catálogo, Plantillas, Productos físicos y Ayuda. En dispositivos pequeños se abre desde el botón de menú.")
    agregar_subtitulo(documento, "Buscar stickers")
    agregar_lista(documento, [
        "Abra Catálogo para ver los stickers publicados.",
        "Use el buscador para escribir una palabra del nombre, descripción o SKU.",
        "Elija una categoría si desea limitar los resultados por tema.",
        "Use Anterior y Siguiente cuando el catálogo tenga más de una página.",
    ])
    agregar_subtitulo(documento, "Consultar un producto")
    agregar_parrafo(documento, "Haga clic sobre una tarjeta para abrir la ficha del producto. Allí podrá revisar sus imágenes, descripción y precio antes de agregarlo a su selección.")

    agregar_titulo(documento, "3. Armar una selección")
    agregar_lista(documento, [
        "Presione Agregar en una tarjeta de producto o en su ficha.",
        "Abra Mi selección desde el botón ubicado en el encabezado.",
        "Aumente o reduzca la cantidad con los controles + y −.",
        "Use × para quitar un producto o Vaciar selección para empezar nuevamente.",
    ])
    agregar_nota(documento, "La selección se conserva solo en el navegador que está usando. No crea una cuenta de cliente ni queda guardada en una base de datos de Papelería de Sol.")

    agregar_titulo(documento, "4. Continuar por WhatsApp")
    agregar_parrafo(documento, "Cuando la selección esté lista, presione Continuar por WhatsApp. La tienda mostrará un aviso antes de abrir la aplicación o la versión web de WhatsApp.")
    agregar_lista(documento, [
        "Lea el aviso y, si está de acuerdo, seleccione Abrir WhatsApp.",
        "Revise el mensaje precompletado antes de enviarlo. Incluye los productos, cantidades y el total mostrado en la tienda.",
        "Consulte por disponibilidad, precio final, pago, retiro o envío directamente en la conversación.",
    ])
    agregar_nota(documento, "La tienda no procesa pagos. En la etapa actual, los medios de pago y la entrega se acuerdan por WhatsApp.")

    agregar_titulo(documento, "5. Privacidad y condiciones")
    agregar_parrafo(documento, "Desde el pie de página y desde el aviso previo a WhatsApp puede abrir la Política de privacidad y los Términos y condiciones. Estos textos se muestran en una ventana dentro de la tienda y también cuentan con una página pública propia.")
    agregar_lista(documento, [
        "Política de privacidad: explica qué ocurre con la selección y con los datos que usted comparta voluntariamente por WhatsApp.",
        "Términos y condiciones: describe la modalidad actual de consulta, pago, entrega, productos personalizados, plazos y atención de fallas.",
    ])

    agregar_titulo(documento, "6. Ayuda y resolución de problemas")
    agregar_subtitulo(documento, "No veo un producto")
    agregar_parrafo(documento, "Pruebe con otra búsqueda o categoría. Si sigue sin encontrarlo, use la página Ayuda o cualquier botón de WhatsApp para consultar.")
    agregar_subtitulo(documento, "La selección quedó vacía")
    agregar_parrafo(documento, "La selección se guarda en el navegador. Si se borran los datos del navegador, se usa otro dispositivo o se navega en modo privado, es posible que deba agregar los productos nuevamente.")
    agregar_subtitulo(documento, "Necesito una respuesta")
    agregar_parrafo(documento, "Todas las consultas comerciales se atienden por WhatsApp. Incluya el nombre del producto y cualquier detalle importante para recibir una respuesta más clara.")

    return documento


def crear_manual_administracion():
    documento = configurar_documento("Manual para administración")
    agregar_portada(
        documento,
        "Manual de uso para administración",
        "Guía para actualizar productos, categorías, contenidos y apariencia de la tienda.",
    )

    agregar_titulo(documento, "1. Alcance y acceso seguro")
    agregar_parrafo(documento, "El dashboard es exclusivo para administradores autorizados. Se accede desde /admin con el correo y la contraseña de administración.")
    agregar_lista(documento, [
        "No comparta la contraseña ni deje abierta la sesión en equipos compartidos.",
        "La sesión expira después de 30 minutos sin actividad.",
        "Use Cerrar sesión al terminar. La tienda solicitará confirmación antes de cerrar la sesión manualmente.",
    ])
    agregar_nota(documento, "No publique un producto hasta revisar sus datos e imágenes. La opción No publicado permite preparar el contenido sin que aparezca en la tienda.")

    agregar_titulo(documento, "2. Resumen del dashboard")
    agregar_parrafo(documento, "Al ingresar verá un resumen de productos, publicaciones, categorías, temas y secciones. Use el menú lateral para abrir Productos, Categorías, Temas de stickers, Secciones del sitio, Datos comerciales, Contenido del sitio o Ayuda.")

    agregar_titulo(documento, "3. Crear y editar productos")
    agregar_subtitulo(documento, "Antes de crear")
    agregar_parrafo(documento, "Si el producto pertenece a un grupo nuevo, cree primero una categoría del tipo adecuado. Los tipos disponibles son Sticker, Plantilla y Producto físico.")
    agregar_subtitulo(documento, "Completar el formulario")
    agregar_lista(documento, [
        "Seleccione Crear nuevo producto y escriba el nombre.",
        "Elija el tipo de producto. Esto habilita las categorías correspondientes y permite sugerir el SKU.",
        "Use el icono ! junto a SKU para leer la ayuda. La sugerencia se completa al elegir el tipo, pero puede ajustarse si es necesario.",
        "Escriba una descripción clara y un precio en pesos. Ambos datos son obligatorios.",
        "Elija No publicado para guardar el producto sin mostrarlo o Publicado para que aparezca al público.",
        "Active Controlar stock solo si desea indicar unidades disponibles.",
        "Marque Mostrar como destacado solo cuando deba aparecer en los espacios destacados de la tienda.",
        "Cargue hasta cinco imágenes. Cada archivo debe pesar como máximo 5 MB.",
    ])
    agregar_subtitulo(documento, "Opciones avanzadas")
    agregar_parrafo(documento, "El desplegable Opciones avanzadas contiene el título y la descripción para buscadores. Déjelos vacíos para usar el nombre y la descripción del producto; complételos únicamente cuando necesite una versión específica para resultados de búsqueda.")
    agregar_subtitulo(documento, "Modificar, archivar o restaurar")
    agregar_parrafo(documento, "Desde el listado seleccione Editar para corregir un registro. En productos, Archivar lo retira del listado activo sin borrarlo; el filtro de archivados permite restaurarlo como No publicado cuando vuelva a ser necesario.")

    agregar_titulo(documento, "4. Administrar categorías y temas")
    agregar_parrafo(documento, "Las categorías ordenan productos dentro de cada tipo. Para stickers pueden representar temas como fútbol, memes o Star Wars.")
    agregar_lista(documento, [
        "Cree una categoría con su nombre y tipo de producto correcto.",
        "Agregue una descripción si ayuda a identificar su uso interno o público.",
        "Marque Publicada cuando deba estar disponible en los filtros o secciones públicas.",
        "Use Orden para decidir la prioridad de aparición cuando exista más de una categoría del mismo tipo.",
    ])

    agregar_titulo(documento, "5. Editar secciones y contenidos")
    agregar_subtitulo(documento, "Secciones del sitio")
    agregar_parrafo(documento, "Aquí se modifican las secciones existentes de la tienda. Se pueden actualizar título, subtítulo, contenido, texto y enlace de botón, publicación y datos SEO. No se crean secciones nuevas desde este lugar cuando no existe un bloque visual que las utilice.")
    agregar_subtitulo(documento, "Apariencia independiente")
    agregar_parrafo(documento, "Dentro de una sección, Apariencia de cada texto permite cambiar color, fuente, tamaño, negrita y cursiva de título, subtítulo y contenido de forma independiente. Estos ajustes afectan solo esa sección.")
    agregar_subtitulo(documento, "Contenido del sitio")
    agregar_parrafo(documento, "Los desplegables de Contenido del sitio reúnen los textos del encabezado, inicio, catálogo, páginas de productos, selección, WhatsApp y Ayuda. Guarde cada bloque después de modificarlo.")

    agregar_titulo(documento, "6. Datos comerciales e identidad visual")
    agregar_parrafo(documento, "En Datos comerciales se actualizan los datos que aparecen en la tienda, como WhatsApp, correo, redes y referencias de atención. Revise cuidadosamente los enlaces antes de guardar.")
    agregar_parrafo(documento, "En Paleta y tipografía general se define la identidad visual completa. Puede cambiar la paleta y elegir una fuente general. Estas opciones afectan a toda la tienda; los estilos de una sección son más específicos.")
    agregar_nota(documento, "Después de guardar, actualice una pestaña pública de la tienda para comprobar el resultado. Pruebe siempre también en un teléfono o con el navegador en tamaño reducido.")

    agregar_titulo(documento, "7. Flujo recomendado antes de publicar")
    agregar_lista(documento, [
        "Cree o edite el contenido como No publicado.",
        "Revise nombre, precio, descripción, categoría, SKU e imágenes.",
        "Abra la vista pública en otra pestaña y verifique cómo se muestra en computadora y celular.",
        "Cambie el estado a Publicado cuando todo esté listo.",
        "Pruebe el botón de WhatsApp para confirmar que el mensaje y el número son correctos.",
    ])

    agregar_titulo(documento, "8. Problemas frecuentes")
    agregar_subtitulo(documento, "No puedo iniciar sesión")
    agregar_parrafo(documento, "Verifique el correo y contraseña. Si el mensaje indica falta de permisos, confirme que esa cuenta fue dada de alta como administradora en Supabase. No cree cuentas públicas para administrar la tienda.")
    agregar_subtitulo(documento, "No se guardan los cambios")
    agregar_parrafo(documento, "Compruebe que la sesión siga activa y que los campos obligatorios estén completos. Si la sesión venció, vuelva a iniciar sesión y repita el cambio.")
    agregar_subtitulo(documento, "No aparece una imagen o producto")
    agregar_parrafo(documento, "Revise que la carga de imagen haya terminado y que el producto esté en estado Publicado. Si corresponde a una categoría, confirme también que la categoría esté publicada.")

    return documento


def guardar_documento(documento, nombre):
    CARPETA_SALIDA.mkdir(parents=True, exist_ok=True)
    destino = CARPETA_SALIDA / nombre
    documento.core_properties.title = nombre.replace("-", " ").replace(".docx", "").title()
    documento.core_properties.author = "Papelería de Sol"
    documento.core_properties.subject = "Manual de uso"
    documento.core_properties.comments = "Documento generado para la operación de Papelería de Sol."
    documento.save(destino)
    print(destino)


if __name__ == "__main__":
    guardar_documento(crear_manual_clientes(), "manual-uso-clientes.docx")
    guardar_documento(crear_manual_administracion(), "manual-uso-administracion.docx")
